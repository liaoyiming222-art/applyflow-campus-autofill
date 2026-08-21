from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "虚构测试简历-林晨曦.docx"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 98, 112)
LIGHT = "E8EEF5"


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_labeled_paragraph(doc, label, value, after=3):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    set_font(paragraph.add_run(f"{label}："), bold=True, color=DARK)
    set_font(paragraph.add_run(value))
    return paragraph


def add_record_marker(doc, number):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run(f"第 {number} 段"), size=10, bold=True, color=MUTED)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 8),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
):
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# customer_pack-style left-aligned resume masthead
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
set_font(title.add_run("林晨曦"), size=26, bold=True, color=DARK)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
set_font(subtitle.add_run("软件开发工程师 · 2027 届校招"), size=13, bold=True, color=BLUE)
set_font(subtitle.add_run("   |   虚构测试简历"), size=10, color=MUTED)

doc.add_heading("基本信息", level=1)
for label, value in (
    ("姓名", "林晨曦"), ("英文姓名", "Chenxi Lin"),
    ("手机号", "13800138027"), ("邮箱", "chenxi.lin.test@example.com"),
    ("性别", "女"), ("出生日期", "2003-08-16"),
    ("当前城市", "上海"), ("籍贯", "江苏省苏州市"), ("民族", "汉族"),
    ("国籍", "中国"), ("证件类型", "身份证"), ("证件号码", "TEST-CERT-2027-001"),
    ("政治面貌", "共青团员"), ("婚姻状况", "未婚"), ("身高", "165"), ("体重", "52"),
    ("微信ID", "chenxi_lin_demo"),
    ("个人主页", "https://github.com/chenxi-lin-demo"),
):
    add_labeled_paragraph(doc, label, value, after=2)

doc.add_heading("教育经历", level=1)
add_record_marker(doc, 1)
for label, value in (
    ("学校", "上海交通大学"), ("学院", "计算机学院"), ("所在校区", "闵行校区"),
    ("专业", "计算机科学与技术"), ("学历", "硕士研究生"), ("学位", "工学硕士"),
    ("入学时间", "2024-09"), ("毕业时间", "2027-06"),
    ("GPA", "3.82/4.00"), ("学习成绩排名", "前 15%"),
    ("是否海外院校", "否"), ("学历类型", "全日制"), ("一级学科", "计算机科学与技术"),
    ("主修课程", "高级算法、分布式系统、机器学习、软件工程方法"),
    ("实验室", "智能软件工程实验室"), ("领域方向", "智能软件工程与人机协作"),
    ("导师", "周明远（虚构）"),
    ("是否学生干部", "是"), ("学生干部名称", "研究生会技术部部长"), ("是否最高学历", "是"),
):
    add_labeled_paragraph(doc, label, value)

add_record_marker(doc, 2)
for label, value in (
    ("学校", "苏州大学"), ("学院", "计算机科学与技术学院（软件学院）"), ("所在校区", "北校区"),
    ("专业", "软件工程"), ("学历", "本科"), ("学位", "工学学士"),
    ("入学时间", "2020-09"), ("毕业时间", "2024-06"),
    ("GPA", "3.72/4.00"), ("学习成绩排名", "前 15%"),
    ("是否海外院校", "否"), ("学历类型", "全日制"), ("一级学科", "计算机科学与技术"),
    ("主修课程", "数据结构、操作系统、计算机网络、数据库系统、软件工程"),
    ("实验室", "软件工程实验室"), ("领域方向", "Web 前端工程与软件质量"),
    ("导师", "陈知行（虚构）"),
    ("是否学生干部", "否"), ("是否最高学历", "否"),
):
    if value:
        add_labeled_paragraph(doc, label, value)

doc.add_heading("实习经历", level=1)
add_record_marker(doc, 1)
for label, value in (
    ("公司名称", "星河互联科技有限公司"), ("部门", "企业效率产品部"),
    ("职位", "前端开发实习生"), ("开始时间", "2026-06"), ("结束时间", "2026-09"),
    ("工作内容", "参与企业协同平台的表单引擎开发，负责字段配置、校验提示与数据回显；使用 TypeScript 和 Vue 3 完成 12 个通用组件，并补充 40 余条单元测试。"),
):
    add_labeled_paragraph(doc, label, value)

doc.add_page_break()
doc.add_heading("实习经历", level=1)
add_record_marker(doc, 2)
for label, value in (
    ("公司名称", "云帆数据实验室"), ("部门", "数据应用组"),
    ("职位", "软件开发实习生"), ("开始时间", "2025-07"), ("结束时间", "2025-10"),
    ("工作内容", "协助开发日志分析工具，完成数据清洗、异常聚合和可视化页面；将常用查询平均响应时间从 2.1 秒优化至 0.8 秒。"),
):
    add_labeled_paragraph(doc, label, value)

doc.add_heading("项目经历", level=1)
add_record_marker(doc, 1)
for label, value in (
    ("项目名称", "校园活动智能管理平台"), ("项目角色", "项目负责人 / 全栈开发"),
    ("开始时间", "2025-09"), ("结束时间", "2026-03"),
    ("项目描述", "面向学生社团的活动发布、报名、签到与统计平台。负责需求拆分、数据库设计和核心模块开发，使用 React、Node.js 与 PostgreSQL，实现二维码签到和活动数据看板。"),
    ("项目成果", "在 6 个学生组织中试运行，累计管理 38 场活动、1,200 余次报名；将人工汇总时间由每场约 30 分钟降至 5 分钟以内。"),
):
    add_labeled_paragraph(doc, label, value)

add_record_marker(doc, 2)
for label, value in (
    ("项目名称", "轻量级浏览器表单助手"), ("项目角色", "独立开发者"),
    ("开始时间", "2026-02"), ("结束时间", "2026-05"),
    ("项目描述", "基于 Manifest V3 开发浏览器扩展，研究网页字段语义识别、表单事件触发与本地数据存储，完成文本框、日期和原生下拉控件的自动填写原型。"),
    ("项目成果", "构建 80 余个字段别名测试样例，在自建表单中实现常用字段稳定识别与填写。"),
):
    add_labeled_paragraph(doc, label, value)

doc.add_heading("专业技能", level=1)
add_labeled_paragraph(doc, "技能", "JavaScript、TypeScript、HTML、CSS、Vue 3、React、Node.js、Python、SQL、Git、基础单元测试与浏览器扩展开发")

doc.add_heading("外语能力", level=1)
add_record_marker(doc, 1)
for label, value in (
    ("语言", "英语"), ("考试类型", "大学英语"), ("语言等级", "六级"),
    ("等级考试得分", "523"), ("熟练程度", "熟练"),
    ("补充说明", "可阅读英文技术文档并进行日常书面沟通"),
):
    add_labeled_paragraph(doc, label, value)

doc.add_heading("奖励荣誉", level=1)
add_record_marker(doc, 1)
for label, value in (
    ("奖励级别", "省级"), ("奖励名称", "全国大学生计算机设计大赛华东赛区二等奖"),
    ("获奖时间", "2025-08"),
    ("详细描述", "担任项目主要开发成员，负责核心功能实现与现场答辩，获得华东赛区二等奖。"),
):
    add_labeled_paragraph(doc, label, value)
add_record_marker(doc, 2)
for label, value in (
    ("奖励级别", "校级"), ("奖励名称", "2024-2025 学年校级二等奖学金"),
    ("获奖时间", "2025-06"),
    ("详细描述", "依据学年综合成绩与实践表现评定，获得校级二等奖学金。"),
):
    add_labeled_paragraph(doc, label, value)

doc.add_heading("求职偏好", level=1)
for label, value in (("期望岗位", "前端开发工程师 / 软件开发工程师"), ("期望城市", "上海、杭州、苏州"), ("可到岗日期", "2027-07-01"), ("招聘信息来源", "校园招聘官网"), ("期望薪酬", "15000")):
    add_labeled_paragraph(doc, label, value)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("ApplyFlow 虚构测试简历 · 姓名、联系方式及公司均为虚构；学校与学院名称用于真实下拉选项测试"), size=8.5, color=MUTED)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
